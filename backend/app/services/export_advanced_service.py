from typing import Dict, List, Optional
import os
import uuid
from datetime import datetime
from app.core.config import settings
from app.services.story_storage import StoryStorage
try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None
try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    markdown = None


class ExportAdvancedService:
    def __init__(self):
        self.story_storage = StoryStorage()
        self.export_path = os.path.join(settings.STORAGE_PATH, "exports")
        os.makedirs(self.export_path, exist_ok=True)
    
    def export_to_word(
        self,
        story_id: str,
        include_images: bool = False
    ) -> Dict:
        """Word formatına dışa aktarır."""
        if not Document:
            raise ValueError("python-docx kütüphanesi yüklü değil")
        
        story = self.story_storage.get_story(story_id)
        if not story:
            raise ValueError("Hikâye bulunamadı")
        
        doc = Document()
        doc.add_heading(story.get('theme', 'Hikâye'), 0)
        
        if story.get('story_text'):
            doc.add_paragraph(story.get('story_text'))
        
        if include_images and story.get('image_url'):
            # Görsel ekleme (gerçek implementasyon için görsel indirme gerekli)
            pass
        
        file_path = os.path.join(self.export_path, f"{story_id}.docx")
        doc.save(file_path)
        
        return {
            "story_id": story_id,
            "format": "docx",
            "file_path": file_path,
            "exported_at": datetime.now().isoformat()
        }
    
    def export_to_markdown(
        self,
        story_id: str
    ) -> Dict:
        """Markdown formatına dışa aktarır."""
        story = self.story_storage.get_story(story_id)
        if not story:
            raise ValueError("Hikâye bulunamadı")
        
        md_content = f"# {story.get('theme', 'Hikâye')}\n\n"
        md_content += f"{story.get('story_text', '')}\n"
        
        file_path = os.path.join(self.export_path, f"{story_id}.md")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return {
            "story_id": story_id,
            "format": "markdown",
            "file_path": file_path,
            "exported_at": datetime.now().isoformat()
        }
    
    def export_to_html(
        self,
        story_id: str,
        include_styles: bool = True
    ) -> Dict:
        """HTML formatına dışa aktarır."""
        story = self.story_storage.get_story(story_id)
        if not story:
            raise ValueError("Hikâye bulunamadı")
        
        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{story.get('theme', 'Hikâye')}</title>
"""
        if include_styles:
            html_content += """
    <style>
        body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }
        h1 { color: #333; }
        p { line-height: 1.6; }
    </style>
"""
        html_content += """
</head>
<body>
"""
        html_content += f"<h1>{story.get('theme', 'Hikâye')}</h1>\n"
        html_content += f"<p>{story.get('story_text', '').replace(chr(10), '</p><p>')}</p>\n"
        html_content += """
</body>
</html>
"""
        
        file_path = os.path.join(self.export_path, f"{story_id}.html")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return {
            "story_id": story_id,
            "format": "html",
            "file_path": file_path,
            "exported_at": datetime.now().isoformat()
        }
    
    def batch_export(
        self,
        story_ids: List[str],
        format: str = "markdown"
    ) -> Dict:
        """Toplu dışa aktarma."""
        results = []
        for story_id in story_ids:
            try:
                if format == "word":
                    result = self.export_to_word(story_id)
                elif format == "html":
                    result = self.export_to_html(story_id)
                else:
                    result = self.export_to_markdown(story_id)
                results.append(result)
            except Exception as e:
                results.append({"story_id": story_id, "error": str(e)})
        
        return {
            "total": len(story_ids),
            "successful": len([r for r in results if "error" not in r]),
            "failed": len([r for r in results if "error" in r]),
            "results": results
        }

