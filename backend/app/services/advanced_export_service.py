from typing import Dict, Optional
import os
import uuid
from datetime import datetime
from app.core.config import settings
from app.services.story_storage import StoryStorage
try:
    from docx import Document
except ImportError:
    Document = None


class AdvancedExportService:
    def __init__(self):
        self.story_storage = StoryStorage()
        self.exports_path = os.path.join(settings.STORAGE_PATH, "exports")
        self._ensure_directory()
    
    def _ensure_directory(self):
        os.makedirs(self.exports_path, exist_ok=True)
    
    def export_to_word(self, story_id: str, title: Optional[str] = None) -> Dict:
        if Document is None:
            raise ValueError("python-docx paketi yüklü değil")
        
        story = self.story_storage.get_story(story_id)
        if not story:
            raise ValueError("Hikâye bulunamadı")
        
        doc = Document()
        doc.add_heading(title or story.get('theme', 'Hikâye'), 0)
        doc.add_paragraph(story.get('story_text', ''))
        
        export_id = str(uuid.uuid4())
        file_path = os.path.join(self.exports_path, f"{export_id}.docx")
        doc.save(file_path)
        
        return {
            "export_id": export_id,
            "file_path": f"/storage/exports/{export_id}.docx",
            "format": "docx",
            "created_at": datetime.now().isoformat()
        }
    
    def export_to_markdown(self, story_id: str) -> Dict:
        story = self.story_storage.get_story(story_id)
        if not story:
            raise ValueError("Hikâye bulunamadı")
        
        md_content = f"# {story.get('theme', 'Hikâye')}\n\n{story.get('story_text', '')}"
        
        export_id = str(uuid.uuid4())
        file_path = os.path.join(self.exports_path, f"{export_id}.md")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return {
            "export_id": export_id,
            "file_path": f"/storage/exports/{export_id}.md",
            "format": "markdown",
            "created_at": datetime.now().isoformat()
        }
    
    def export_to_html(self, story_id: str) -> Dict:
        story = self.story_storage.get_story(story_id)
        if not story:
            raise ValueError("Hikâye bulunamadı")
        
        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <title>{story.get('theme', 'Hikâye')}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #333; }}
        p {{ line-height: 1.6; }}
    </style>
</head>
<body>
    <h1>{story.get('theme', 'Hikâye')}</h1>
    <p>{story.get('story_text', '').replace(chr(10), '</p><p>')}</p>
</body>
</html>
"""
        
        export_id = str(uuid.uuid4())
        file_path = os.path.join(self.exports_path, f"{export_id}.html")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return {
            "export_id": export_id,
            "file_path": f"/storage/exports/{export_id}.html",
            "format": "html",
            "created_at": datetime.now().isoformat()
        }

